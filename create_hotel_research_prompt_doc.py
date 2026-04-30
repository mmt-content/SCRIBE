from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Hotel_Research_Tool_Product_Tech_Prompt.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="DDE3EE", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_labelled_bullets(doc, label, items):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    for item in items:
        add_bullet(doc, item)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(30, 41, 59)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 22, "111827"),
        ("Heading 1", 16, "111827"),
        ("Heading 2", 13, "243B53"),
        ("Heading 3", 11.5, "374151"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "Hotel Research Tool - Product + Tech Prompt"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Prepared for MakeMyTrip Scribe"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F7FF")
    set_cell_border(cell, "C7D2FE")
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(67, 56, 202)
    p.add_run("\n" + text)


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "Requirement"
    for cell in hdr:
        set_cell_shading(cell, "EEF2FF")
        set_cell_border(cell)
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for area, req in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = req
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)
            set_cell_margins(cell)
    return table


def build():
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Hotel Research Tool - Product + Tech Prompt")
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Build specification for a Gemini-powered, official-website-only hotel research tool."
    ).italic = True

    add_callout(
        doc,
        "Core Objective",
        "Build a web-based hotel research tool that crawls official hotel websites and returns structured, high-accuracy research output for hotel content and factual research workflows. Accuracy is critical: no assumptions, no hallucinations, and no paraphrased facts during extraction.",
    )

    add_h2(doc, "Tech Requirements")
    add_bullet(doc, "Host on GitHub.")
    add_bullet(doc, "Use Gemini API for all AI processing.")
    add_bullet(doc, "Strictly separate core functionality from the UI/design layer.")
    add_bullet(doc, "Core functionality includes logic, crawling, extraction, AI processing, category handling, and prompt handling.")
    add_bullet(doc, "UI/design layer includes layout, styling, visual states, and interface components.")
    add_bullet(doc, "UI should be easily replaceable without affecting backend logic.")

    add_h2(doc, "1. Input + Research Flow")
    add_bullet(doc, "Input single hotel name or multiple hotel names in batch mode.")
    add_bullet(doc, "Hotel names should support comma-separated input.")
    add_bullet(doc, "User selects MMT Content Mode or Raw Research Mode.")
    add_bullet(doc, "In MMT Content Mode, user selects one content output type.")
    add_bullet(doc, "In Raw Research Mode, user selects one or more data categories.")
    add_bullet(doc, "Primary CTA: Run Research.")

    add_h2(doc, "2. Research Modes + Categories")
    add_h3(doc, "A. MMT Content Mode")
    doc.add_paragraph(
        "This mode generates ready-to-use structured content and marketing assets based on official hotel website research and SOP/category rules."
    )
    add_bullet(doc, "Brown Texts")
    add_bullet(doc, "Experience Cards")
    add_bullet(doc, "Property Descriptions")
    add_bullet(doc, "Additional content output types may be added later.")
    add_bullet(doc, "Each output type should have editable category definitions, prompt rules, output schemas, examples, do/don't rules, and evidence requirements.")

    add_h3(doc, "B. Raw Research Mode")
    doc.add_paragraph(
        "This mode extracts comprehensive factual hotel data from the official website without converting it into final marketing copy."
    )
    for item in [
        "Amenities",
        "Room Types & Descriptions",
        "Dining & Restaurants",
        "Location & Nearby Attractions",
        "Hotel Overview / About",
        "Policies",
        "Images & Media",
        "Reviews & Ratings, only if available on the official hotel website",
        "All Categories",
    ]:
        add_bullet(doc, item)
    add_bullet(doc, "More raw research categories should be addable later.")
    add_bullet(doc, "Raw Research Mode must preserve source wording exactly and return factual, structured, cited research.")

    add_h2(doc, "3. Research Prompt Control")
    for item in [
        "A visible, editable Hotel Research Prompt block must appear in the UI.",
        "The prompt defines what data to extract, how to structure output, which mode/category/SOP rules to follow, and what evidence requirements apply.",
        "Users must be able to modify the prompt anytime before running research.",
        "Output should dynamically change based on prompt edits.",
        "Prompt should be pre-filled with recommended research guidelines for accuracy and consistency.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "4. Research Process")
    for item in [
        "Identify and validate the official hotel website.",
        "Crawl the full official website, including all accessible internal pages.",
        "Extract all relevant information exactly as written on the official website.",
        "Do not rewrite, paraphrase, synonymize, summarize, infer, or embellish extracted facts during extraction.",
        "Preserve original wording, names, labels, amenities, room names, dining outlet names, policies, and experience descriptions exactly as found.",
        "Pass extracted source data through Gemini only for structuring, classification, formatting, and category-specific content generation based on SOP guidelines, selected mode, selected category, and user-edited research prompt.",
        "Return final structured output with citations and source excerpts.",
    ]:
        add_number(doc, item)

    add_labelled_bullets(
        doc,
        "Extraction rules",
        [
            "Crawl all accessible internal pages from the official website.",
            "Ignore third-party OTAs unless explicitly enabled later.",
            "Extract facts exactly as written on the website.",
            "No synonyms or alternative phrasings for factual information.",
            "No hallucinated content.",
            "No inferred amenities, services, locations, policies, room features, or experiences.",
            "If information is not found, mark it as not_found.",
            "Each extracted item must include the source page URL and source excerpt.",
        ],
    )

    add_h2(doc, "5. Output States")
    add_bullet(doc, "Each research job should show In Progress, Completed / Finished, and Failed.")
    add_bullet(doc, "Batch runs should display names like Hotel Batch 2 - [First hotel name].")
    add_bullet(doc, "Same layout logic should apply across all states.")

    add_h2(doc, "6. Content Category System")
    add_two_col_table(
        doc,
        [
            ("MMT Content Mode", "Brown Texts; Experience Cards; Property Descriptions."),
            ("Raw Research Mode", "Amenities; Room Types & Descriptions; Dining & Restaurants; Location & Nearby Attractions; Hotel Overview / About; Policies; Images & Media; Reviews & Ratings if available on official website; All Categories."),
            ("Category management", "Add, edit, activate/deactivate, reuse categories, and keep MMT Content Mode categories separate from Raw Research Mode categories."),
            ("Category definition", "Category name, mode, extraction requirements, output structure, writing style where applicable, do/don't rules, evidence format, missing-information behavior, and JSON output schema."),
        ],
    )

    add_h2(doc, "7. SOP Upload + Frontend Category Builder")
    add_bullet(doc, "Frontend must include SOP/category management where users upload SOP files and create content categories directly in the UI.")
    add_bullet(doc, "Supported formats: PDF, DOCX, TXT, Markdown.")
    for item in [
        "User uploads SOP file.",
        "Tool extracts text from the SOP.",
        "Gemini analyzes the SOP and creates a proposed category definition.",
        "User reviews and edits the proposed category before saving.",
        "User assigns the category to MMT Content Mode or Raw Research Mode.",
        "Saved category becomes available in the correct mode/category selector.",
        "The category automatically generates the editable Hotel Research Prompt.",
    ]:
        add_number(doc, item)
    add_labelled_bullets(
        doc,
        "Category builder fields",
        [
            "Category name",
            "Category mode",
            "Category description",
            "Required research fields",
            "Output format",
            "Writing style",
            "Examples",
            "Do rules",
            "Don't rules",
            "Source/evidence requirements",
            "Missing-information rules",
            "JSON output schema",
        ],
    )
    add_bullet(doc, "SOP uploads should not permanently train the model.")
    add_bullet(doc, "SOPs should be converted into structured, editable category definitions.")
    add_bullet(doc, "Users must approve SOP-derived categories before they are used.")
    add_bullet(doc, "Existing categories should remain editable from the frontend.")

    add_h2(doc, "8. UI Requirements")
    add_bullet(doc, "UI must match the provided reference image exactly.")
    add_labelled_bullets(
        doc,
        "Header should include",
        [
            "MMT Hotel Research / Content Intelligence",
            "Official Website Only indicator",
            "Manage Categories button",
            "User/account area",
        ],
    )
    add_bullet(doc, "Main title: MakeMyTrip - Scribe (Powered by Gemini).")
    add_labelled_bullets(
        doc,
        "Primary screen should include",
        [
            "Hotel names input",
            "Add button",
            "Choose Research Mode section",
            "MMT Content Mode card",
            "Raw Research Mode card",
            "Output type selector for MMT Content Mode",
            "Data category selector for Raw Research Mode",
            "Run Research button",
            "Research status panel showing In Progress, Finished, and Failed counts",
            "Hotel Research Prompt block",
            "Quick Start Guide",
            "Recent Searches",
        ],
    )
    add_bullet(doc, "Two main states: initial screen before research and finished screen/results view.")
    add_bullet(doc, "Same layout logic applies to In Progress and Failed states.")
    add_bullet(doc, "UI should render structured data using cards, tables, tabs, fields, status counters, and evidence sections.")

    add_h2(doc, "9. SOP Integration")
    add_bullet(doc, "Tool must read and follow uploaded/provided SOPs for brown text rules, experience card formats, property description rules, writing style, and constraints.")
    add_bullet(doc, "Output must align with these formats without manual cleanup.")
    add_bullet(doc, "SOP-derived rules must be visible, editable, and reusable through content categories.")

    add_h2(doc, "10. Structured API Output Requirement")
    doc.add_paragraph("The API must return structured machine-readable output, not paragraph-only responses. Gemini responses must be requested, parsed, and validated as JSON.")
    sample = """{
  "hotelName": "string",
  "officialWebsite": "string",
  "researchMode": "mmt_content_mode | raw_research_mode",
  "selectedCategories": ["string"],
  "status": "completed | in_progress | failed",
  "summary": {
    "shortAnswer": "string",
    "confidence": "high | medium | low",
    "missingInformation": ["string"]
  },
  "sections": [
    {
      "title": "string",
      "fields": [
        {
          "label": "string",
          "value": "string | not_found",
          "sourceUrl": "string",
          "sourceExcerpt": "string",
          "confidence": "high | medium | low"
        }
      ]
    }
  ],
  "contentOutputs": [
    {
      "type": "brown_text | experience_card | property_description | custom",
      "title": "string",
      "copy": "string",
      "supportingEvidence": [
        {
          "sourceUrl": "string",
          "sourceExcerpt": "string"
        }
      ],
      "warnings": ["string"]
    }
  ],
  "rawEvidence": [
    {
      "pageTitle": "string",
      "url": "string",
      "extractedText": "string"
    }
  ],
  "errors": []
}"""
    p = doc.add_paragraph()
    run = p.add_run(sample)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8.5)
    set_para_shading(p, "F8FAFC")

    add_labelled_bullets(
        doc,
        "Structured output rules",
        [
            "No unstructured paragraph-only API responses.",
            "Each category must define its own JSON output schema.",
            "API must validate Gemini output before showing it in the UI.",
            "If Gemini returns invalid JSON, the system should retry or mark the job as failed.",
            "Any unsupported claim must be excluded or marked as not_found.",
            "Every generated content output must be linked to source evidence.",
            "UI must render results from structured data such as cards, tables, tabs, and fields.",
        ],
    )

    add_h2(doc, "11. Status Listing Pages")
    doc.add_paragraph(
        "When a user clicks a status card from the home screen, the app should open a dedicated listing page for that status."
    )
    add_bullet(doc, "Required status pages: Finished / Completed, In Progress, and Failed.")
    add_bullet(doc, "Each page should replicate the reference layout, with status-specific data changed.")
    add_labelled_bullets(
        doc,
        "Listing page should include",
        [
            "Header with MakeMyTrip - Scribe (Powered by Gemini), Content Intelligence, Official Website Only badge, Manage Categories button, and user/account area",
            "Back to Home button",
            "Page title based on status: Completed, In Progress, or Failed",
            "Count text such as 2558 processes found",
            "Search bar with placeholder: Search by Process ID, Hotel Name or Hotel ID",
            "Filter button",
            "Results table",
            "Pagination at the bottom",
            "Result count text such as Showing 1 to 8 of 2558 results",
        ],
    )
    add_labelled_bullets(
        doc,
        "Table columns",
        [
            "Process ID",
            "User ID",
            "Hotel / Batch Name",
            "Type",
            "Status",
            "Hotel ID / Batch ID",
            "Completed At / Started At / Failed At",
            "Action",
        ],
    )
    add_bullet(doc, "Type should show pill labels such as Single Hotel and Batch.")
    add_bullet(doc, "Status should show pill labels such as Finished, In Progress, and Failed.")
    add_bullet(doc, "Action should include View Details.")
    add_bullet(doc, "The same page structure should be reused for all three statuses, with only title, status labels, timestamp column, and data changing.")
    add_bullet(doc, "Finished page: title Completed, status pill Finished, timestamp column Completed At.")
    add_bullet(doc, "In Progress page: title In Progress, status pill In Progress, timestamp column Started At.")
    add_bullet(doc, "Failed page: title Failed, status pill Failed, timestamp column Failed At.")
    add_bullet(doc, "Clicking View Details should open the structured research result for that process.")

    add_h2(doc, "Non-Negotiables")
    for item in [
        "Zero hallucination tolerance.",
        "Only extract from official sources.",
        "Crawl all accessible internal pages on the official website.",
        "Extract information exactly as written on the official website.",
        "No inferred, assumed, paraphrased, synonymized, or embellished data during extraction.",
        "High consistency across batch runs.",
        "Every output must include source evidence or explicitly say information was not found.",
        "MMT Content Mode and Raw Research Mode categories must remain separate.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "Future-Proofing")
    for item in [
        "UI and backend must be modular.",
        "Prompt system must be flexible.",
        "Category system must be expandable.",
        "Batch processing must scale.",
        "SOP-derived category definitions must be reusable and editable.",
        "Design changes must not affect crawling, extraction, AI processing, or category logic.",
        "More MMT Content Mode output types and Raw Research Mode data categories must be addable later.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)


def set_para_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


if __name__ == "__main__":
    build()
